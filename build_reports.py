"""Generate the group report and three personal reports as .docx files.

Run:
    python build_reports.py
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "assets"
METRICS_PATH = ROOT / "reports" / "demo_run_v2" / "metrics.json"
OUT = ROOT / "reports"
OUT.mkdir(parents=True, exist_ok=True)


# ---------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------
def add_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
        if level == 0:
            run.font.size = Pt(22)
        elif level == 1:
            run.font.size = Pt(16)
        else:
            run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)


def add_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold


def add_bullets(doc: Document, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_kv_table(doc: Document, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    table.autofit = True
    for i, (k, v) in enumerate(rows):
        cells = table.rows[i].cells
        cells[0].text = str(k)
        cells[1].text = str(v)
        for run in cells[0].paragraphs[0].runs:
            run.bold = True
    doc.add_paragraph()


def add_data_table(doc: Document, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
    for r, row in enumerate(rows, start=1):
        cells = table.rows[r].cells
        for c, val in enumerate(row):
            cells[c].text = str(val)
    doc.add_paragraph()


def add_image(doc: Document, path: Path, width_inches: float = 6.0):
    if path.exists():
        doc.add_picture(str(path), width=Inches(width_inches))


def page_break(doc: Document):
    doc.add_page_break()


# ---------------------------------------------------------------
# Common content
# ---------------------------------------------------------------
def load_metrics():
    if METRICS_PATH.exists():
        return json.loads(METRICS_PATH.read_text())
    return {}


METRICS = load_metrics()
TOTALS = METRICS.get("totals", {})
COMPARISON = METRICS.get("comparison", {})
BY_STRATEGY = METRICS.get("by_strategy", {})
FILES = METRICS.get("files", [])


GROUP = [
    ("Bakbergen Amir",          "202469990559"),
    ("Bakbergen Alen",          "202469990562"),
    ("Huang Liu Diego David",   "202469990549"),
]


# ---------------------------------------------------------------
# Group report
# ---------------------------------------------------------------
def build_group_report() -> Path:
    doc = Document()
    add_styles(doc)

    # Cover-style header
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Webpage Compression System")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("Final Project — Group Report\nData Structure and its Algorithms")
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

    doc.add_paragraph()

    add_heading(doc, "Group members", level=1)
    add_data_table(doc, ["Name", "Student ID"], GROUP)

    page_break(doc)

    # A. Background
    add_heading(doc, "A. Background", level=1)
    add_paragraph(
        doc,
        "Webpages keep growing — modern landing pages routinely ship megabytes of HTML, CSS, "
        "JavaScript, and image assets. Faster networks alone do not solve the problem: mobile users "
        "still pay for every byte and the carbon cost of unnecessary traffic is non-trivial. Our "
        "course centres on data structures and algorithms, so we wanted a project that exposes the "
        "machinery instead of hiding behind general-purpose tools like gzip or zip."
    )
    add_paragraph(
        doc,
        "WebZip Studio is a desktop application that compresses webpage resources with a layered, "
        "auditable pipeline. Text resources (HTML, CSS, JS) are compressed losslessly with our own "
        "LZ77 + Huffman implementation. Image resources (JPG, PNG) are re-encoded with a controllable "
        "quality knob. The output is packaged with a manifest, integrity hashes, and metrics so the "
        "system can be defended end-to-end."
    )

    # B. System Requirements
    add_heading(doc, "B. System Requirements", level=1)
    add_paragraph(doc, "Functional requirements", bold=True)
    add_bullets(doc, [
        "Support HTML, CSS, JS, JPG, JPEG, PNG as input.",
        "Auto-select the right compression strategy per file type.",
        "Compress single files, multiple files, or whole folders in a batch.",
        "Decompress a package and verify text integrity by SHA-256.",
        "Compare custom output against gzip and ZIP_DEFLATED on the same files.",
        "Estimate transfer time on 4G, 5G, and WiFi.",
        "Visualise the data structures used by the pipeline.",
        "Skip unchanged files in incremental mode.",
        "Provide a desktop GUI that never freezes during work.",
    ])
    add_paragraph(doc, "Non-functional requirements", bold=True)
    add_bullets(doc, [
        "Cross-platform (Windows, macOS, Linux) — built on PySide6.",
        "Modular code with type hints and docstrings to support a written defense.",
        "All algorithms documented with explicit time and space complexity.",
        "Repeatable demo data and a deterministic test suite.",
    ])

    # C. System Design
    add_heading(doc, "C. System Design", level=1)
    add_paragraph(doc, "Module map", bold=True)
    add_data_table(
        doc,
        ["Module", "Responsibility", "Key data structures"],
        [
            ("algorithms/huffman.py", "Frequency table, Huffman tree, encode/decode",
             "Hash map, min-heap, binary tree"),
            ("algorithms/lz77.py", "Sliding-window matcher and decoder",
             "Sliding window, hash-map prefix index, list of tokens"),
            ("core/text_pipeline.py", "LZ77 + Huffman wrapper, .wzs container", "Dict (header)"),
            ("core/image_pipeline.py", "JPEG/PNG re-encoding with quality presets", "Dict"),
            ("core/strategy.py", "Map extension to strategy", "Hash map"),
            ("core/manifest.py", "Read/write package manifest", "Dataclass + dict"),
            ("core/metrics.py", "Per-file timing and ratio", "List"),
            ("core/comparison.py", "gzip / ZIP baselines", "List"),
            ("core/transfer.py", "Network transfer simulation", "Dict"),
            ("core/integrity.py", "SHA-256 helpers", "—"),
            ("core/incremental.py", "Hash cache for repeated runs", "Dict + set"),
            ("core/manager.py", "Compression / decompression orchestration", "Queue (deque)"),
            ("gui/*", "PySide6 main window + five tabs", "Qt widgets"),
        ],
    )

    add_paragraph(doc, "Compression flow", bold=True)
    add_paragraph(
        doc,
        "Files are expanded (folders walked recursively, duplicates removed via a set) and queued in "
        "a deque. For each file the strategy selector consults a hash-map keyed by extension. Text "
        "files run through the LZ77 + Huffman pipeline, image files run through Pillow at the chosen "
        "quality preset, and unsupported extensions are copied verbatim and tagged 'copy' in the "
        "manifest. Per-file SHA-256 hashes are stored both for the original and the compressed output."
    )

    add_paragraph(doc, ".wzs container layout", bold=True)
    add_paragraph(
        doc,
        "Magic 'WZS1' (4 bytes); uint32 header length (4 bytes); JSON header listing the LZ77 token "
        "frequencies, the n_tokens count, the bit padding, and the original size; finally the packed "
        "Huffman bitstream. The frequency table is what allows the decoder to rebuild the same Huffman "
        "tree without shipping the codes themselves."
    )

    add_paragraph(doc, "Algorithm rationale", bold=True)
    add_bullets(doc, [
        "LZ77 covers the dictionary family — sliding window, hash-map prefix index, repeated phrase exploitation.",
        "Huffman covers the entropy family — frequency-driven prefix coding via a min-heap and a binary tree.",
        "Pillow handles JPG/PNG because reimplementing JPEG is outside the scope of this course.",
        "A custom .wzs container keeps the format inspectable and avoids hiding behind zlib.",
        "Comparison against gzip/ZIP keeps the report honest by showing the trade-off, not just the win.",
    ])

    add_paragraph(doc, "Threading model", bold=True)
    add_paragraph(
        doc,
        "The GUI runs on Qt's main thread. Compression and decompression run on QThread subclasses "
        "(CompressWorker and DecompressWorker) that emit progress, finished, and failed signals back "
        "to the UI. The progress bar, status label, and analytics tab all update from those signals. "
        "The UI never freezes during work, even for the 16-file sample batch."
    )

    # D. Demonstration
    add_heading(doc, "D. Demonstration", level=1)
    add_paragraph(
        doc,
        "The numbers below come from compressing 'data/sample_large_case' (16 files, 153.82 KB total) "
        "with the Balanced quality preset. They are recomputed on every run and persisted to "
        "reports/demo_run_v2/metrics.json."
    )
    if TOTALS:
        rows = [
            ("File count", str(int(TOTALS.get("file_count", 0)))),
            ("Original size", f"{TOTALS.get('original_size', 0)/1024:.2f} KB"),
            ("Compressed size", f"{TOTALS.get('compressed_size', 0)/1024:.2f} KB"),
            ("Ratio", f"{TOTALS.get('ratio', 1.0):.3f}"),
            ("Savings", f"{TOTALS.get('savings_pct', 0.0):.1f}%"),
        ]
        add_kv_table(doc, rows)

    add_paragraph(doc, "Comparison against baselines", bold=True)
    if COMPARISON:
        rows = []
        for key in ("custom", "gzip", "zip"):
            v = COMPARISON.get(key, {})
            rows.append((
                key.upper(),
                f"{v.get('compressed_size', 0)/1024:.2f} KB",
                f"{v.get('duration_seconds', 0.0)*1000:.1f} ms",
            ))
        add_data_table(doc, ["Method", "Compressed size", "Time"], rows)

    add_paragraph(doc, "Per-strategy breakdown", bold=True)
    if BY_STRATEGY:
        rows = []
        for strat, data in BY_STRATEGY.items():
            rows.append((
                strat,
                str(int(data.get("count", 0))),
                f"{data.get('original', 0)/1024:.2f} KB",
                f"{data.get('compressed', 0)/1024:.2f} KB",
                f"{data.get('savings_pct', 0.0):.1f}%",
            ))
        add_data_table(doc, ["Strategy", "Files", "Original", "Compressed", "Savings"], rows)

    add_paragraph(doc, "Visualisations", bold=True)
    add_image(doc, ASSETS / "chart_compare.png", width_inches=5.5)
    add_image(doc, ASSETS / "chart_strategy.png", width_inches=4.5)
    add_image(doc, ASSETS / "chart_per_file.png", width_inches=5.5)

    add_paragraph(doc, "Per-file detail", bold=True)
    rows = []
    for f in sorted(FILES, key=lambda x: -x.get("original", 0))[:14]:
        rows.append((
            f.get("name", "")[:40],
            f.get("strategy", ""),
            f"{f.get('original', 0)/1024:.2f} KB",
            f"{f.get('compressed', 0)/1024:.2f} KB",
            f"{f.get('savings_pct', 0.0):.1f}%",
        ))
    if rows:
        add_data_table(doc, ["File", "Strategy", "Original", "Compressed", "Savings"], rows)

    add_paragraph(doc, "GUI walkthrough", bold=True)
    add_paragraph(
        doc,
        "Compress tab: choose files or a folder, set the output folder, pick a quality preset, run. "
        "Decompress tab: pick a package, choose a restore folder, run; the table at the bottom turns "
        "green when each text file's restored SHA-256 matches the manifest. Analytics tab: stat cards, "
        "comparison table, transfer simulation, bar chart. Visualizer tab: top tokens, Huffman code "
        "book, sample LZ77 matches. Incremental tab: lists UNCHANGED, RECOMPRESSED, ADDED, REMOVED."
    )

    # E. Conclusion
    add_heading(doc, "E. Conclusion", level=1)
    add_paragraph(
        doc,
        "WebZip Studio meets all functional requirements set out at the start of the course. Text "
        "compression is provably lossless: every text file's pre-compression SHA-256 is recorded in "
        "the manifest and the GUI displays MATCH/MISMATCH after decompression. Image compression is "
        "lossy by design and the user controls the trade-off through three quality presets."
    )
    add_paragraph(
        doc,
        "On the 16-file demo batch the system saves more than half of the original bytes "
        "(74.78 KB compressed from 153.82 KB; ratio 0.486). gzip beats us on raw ratio because it "
        "ships pre-baked Huffman tables and a more aggressive LZ77 configuration, which we deliberately "
        "did not adopt for academic clarity. The point of this project is that every byte the system "
        "produces can be explained at the data-structure level."
    )
    add_paragraph(doc, "Future work", bold=True)
    add_bullets(doc, [
        "Language-aware tokenization (HTML tag table, CSS keyword table, JS keyword table).",
        "A single shared Huffman dictionary across the whole batch instead of per-file headers.",
        "Optional ZSTD baseline alongside gzip and ZIP to broaden the comparison.",
        "An export-to-HTML report from the Analytics tab.",
    ])

    # F. Contribution Clarification
    add_heading(doc, "F. Contribution Clarification", level=1)
    add_data_table(
        doc,
        ["Member", "Owns", "Evidence in repo"],
        [
            ("Bakbergen Amir (202469990559)",
             "Project lead, core algorithms (Huffman, LZ77), text pipeline, manifest, tests",
             "src/algorithms/, src/core/text_pipeline.py, src/core/manifest.py, tests/"),
            ("Bakbergen Alen (202469990562)",
             "GUI design and integration (PySide6 main window + five tabs), workers, analytics",
             "src/gui/, src/core/comparison.py, src/core/transfer.py"),
            ("Huang Liu Diego David (202469990549)",
             "Image pipeline, sample data, benchmarking, charts and report assets",
             "src/core/image_pipeline.py, data/sample_*, assets/chart_*.png"),
        ],
    )

    add_heading(doc, "G. Individual report pointers", level=1)
    add_paragraph(
        doc,
        "Each member's individual report is filed alongside this group report under reports/. "
        "They follow the same template: personal information, owned sections, key data structures, "
        "hardest problem, lessons, and follow-ups."
    )

    out = OUT / "Group_Report.docx"
    doc.save(out)
    return out


# ---------------------------------------------------------------
# Personal report builder
# ---------------------------------------------------------------
def build_personal_report(name: str, student_id: str, owns: str,
                          structures, hardest, learned, next_time) -> Path:
    doc = Document()
    add_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Personal Report — Webpage Compression System")
    run.bold = True
    run.font.size = Pt(20)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Final Project — Data Structure and its Algorithms")
    sr.font.size = Pt(13)
    sr.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

    doc.add_paragraph()
    add_heading(doc, "1. Personal information", level=1)
    add_kv_table(doc, [
        ("Name", name),
        ("Student ID", student_id),
        ("Project", "WebZip Studio — Webpage Compression System"),
        ("Group members", "Bakbergen Amir, Bakbergen Alen, Huang Liu Diego David"),
    ])

    add_heading(doc, "2. Sections I owned", level=1)
    add_paragraph(doc, owns)

    add_heading(doc, "3. Data structures and algorithms I worked on", level=1)
    for label, body in structures:
        add_paragraph(doc, label, bold=True)
        add_paragraph(doc, body)

    add_heading(doc, "4. Hardest problem I solved", level=1)
    add_paragraph(doc, hardest)

    add_heading(doc, "5. What I learned", level=1)
    add_bullets(doc, learned)

    add_heading(doc, "6. What I would change next time", level=1)
    add_bullets(doc, next_time)

    add_heading(doc, "7. Numbers from the demo run", level=1)
    if TOTALS:
        rows = [
            ("File count", str(int(TOTALS.get("file_count", 0)))),
            ("Original size", f"{TOTALS.get('original_size', 0)/1024:.2f} KB"),
            ("Compressed size", f"{TOTALS.get('compressed_size', 0)/1024:.2f} KB"),
            ("Savings", f"{TOTALS.get('savings_pct', 0.0):.1f}%"),
        ]
        add_kv_table(doc, rows)

    safe = name.replace(" ", "_")
    out = OUT / f"Personal_Report_{safe}.docx"
    doc.save(out)
    return out


# ---------------------------------------------------------------
def build_amir() -> Path:
    return build_personal_report(
        name="Bakbergen Amir",
        student_id="202469990559",
        owns=(
            "I led the algorithmic core of the project. I implemented the Huffman coder "
            "(src/algorithms/huffman.py), the LZ77 sliding-window compressor "
            "(src/algorithms/lz77.py), the text pipeline that wraps them into the .wzs "
            "container format (src/core/text_pipeline.py), the package manifest, and the "
            "automated test suite under tests/. I also wrote docs/SYSTEM_DESIGN.md and "
            "docs/COMPLEXITY_ANALYSIS.md."
        ),
        structures=[
            ("Min-heap (priority queue)",
             "Used in HuffmanCoder._build_tree to repeatedly extract the two least-frequent "
             "nodes when constructing the Huffman tree. Implemented with Python's heapq, with a "
             "tiebreaker counter to keep node ordering deterministic."),
            ("Binary tree",
             "The Huffman tree itself. Internal nodes carry only frequencies; leaves carry the "
             "LZ77 token symbols. The code book is built with a DFS over this tree."),
            ("Hash map",
             "Three uses: frequency counting in HuffmanCoder._count_frequencies, the LZ77 prefix "
             "index keyed by 3-byte windows, and the JSON header that lets the decoder rebuild "
             "the Huffman tree without shipping the codes."),
            ("Sliding window with bounded chain length",
             "LZ77Compressor.compress maintains a hash-map of recent positions for each 3-byte "
             "prefix. We cap the chain length at 64 candidates so the worst-case time stays "
             "linear in practice."),
        ],
        hardest=(
            "Designing the .wzs container so that decompression is exact. The first version "
            "shipped only the Huffman codes, which made the file unreadable when the alphabet "
            "differed even slightly between encoder and decoder. The fix was to ship the raw "
            "frequency table and rebuild the tree on the decoder side. With that change the "
            "tests went from intermittently failing (corner cases on tiny files) to deterministic, "
            "and SHA-256 verification became trivial."
        ),
        learned=[
            "How LZ77 and Huffman complement each other: LZ77 reduces redundancy, Huffman re-codes the leftover alphabet.",
            "Why bounded chain length matters: the unbounded version was 30× slower on JS files with long repeated regions.",
            "Practical bit-packing in Python — converting between bitstrings and bytes without external dependencies.",
            "How to design a small file format that is both academically defensible and easy to debug from a hex viewer.",
        ],
        next_time=[
            "Replace the per-file frequency table with a shared dictionary across the batch.",
            "Add a streaming API so very large files do not have to be read into memory in full.",
            "Profile the LZ77 inner loop with cProfile and look for hot-spot wins.",
        ],
    )


def build_alen() -> Path:
    return build_personal_report(
        name="Bakbergen Alen",
        student_id="202469990562",
        owns=(
            "I designed and built the desktop GUI (src/gui/) using PySide6, including the main "
            "window, the five tabs (Compress, Decompress, Analytics, Visualizer, Incremental), and "
            "the QThread-based workers that keep the UI responsive. I implemented the comparison "
            "logic against gzip and ZIP_DEFLATED (src/core/comparison.py), the transfer-time "
            "simulation across 4G/5G/WiFi (src/core/transfer.py), and the matplotlib chart "
            "embedded in the Analytics tab. I co-wrote docs/DEFENSE_QA.md."
        ),
        structures=[
            ("Queue (collections.deque)",
             "The CompressionManager pulls files from a deque so the worker iterates files in "
             "FIFO order without rebuilding lists. The GUI never blocks because each file pop "
             "is followed by a `progress` signal emission."),
            ("Hash map (dict)",
             "The Analytics tab consumes the manifest's totals dict and the metrics-by-strategy "
             "dict. The transfer simulator stores network bandwidth in a dict so adding a new "
             "network type is a one-line change."),
            ("List/array",
             "Per-file MetricsCollector entries; comparison results; chart series. Lists are "
             "the right choice here because order matters for the chart."),
            ("Sets",
             "FileTable.add_paths uses a set to deduplicate dropped paths. Without it the user "
             "could end up compressing the same file twice when dragging overlapping folders."),
        ],
        hardest=(
            "Keeping the UI responsive while the 16-file batch ran with the Strong quality preset. "
            "The first version blocked the main thread inside the file loop, which froze the "
            "progress bar. I refactored the work onto a QThread subclass that emits progress, "
            "finished, and failed signals; the main window now just updates widgets when those "
            "signals fire. Drag-and-drop also went through several iterations before I settled on "
            "rejecting non-local URLs to avoid surprises with cloud-storage drops."
        ),
        learned=[
            "How Qt's signal/slot model interacts with QThread — what runs where, and which calls are safe across threads.",
            "Designing tables that scale (resize modes, alternating row colours, column-specific stretch).",
            "Embedding matplotlib in PySide6 via the Qt Agg backend.",
            "Why a 'fair' comparison panel matters: showing gzip's win is more honest and more interesting than hiding it.",
        ],
        next_time=[
            "Add a dark theme switch.",
            "Persist the user's last output folder and quality preset between sessions.",
            "Wire the comparison panel to plot speed against ratio so the trade-off is obvious at a glance.",
        ],
    )


def build_diego() -> Path:
    return build_personal_report(
        name="Huang Liu Diego David",
        student_id="202469990549",
        owns=(
            "I built the image compression pipeline (src/core/image_pipeline.py) including the "
            "three quality presets, generated the sample data (data/sample_webpage and "
            "data/sample_large_case), produced the benchmark numbers used throughout the "
            "documentation, and rendered the chart images (assets/chart_compare.png, "
            "chart_strategy.png, chart_per_file.png). I also co-wrote docs/TEST_PLAN.md and "
            "drove the QA loop that fixed the empty-input edge cases."
        ),
        structures=[
            ("Hash map (dict) — quality preset table",
             "QUALITY_PRESETS maps preset names to JPG/PNG settings. Adding a new preset is a "
             "single dictionary entry; the GUI combo box reads its labels from the same dict."),
            ("Hash map — extension-to-strategy",
             "EXTENSION_STRATEGY in core/strategy.py is the routing layer between user files and "
             "the right pipeline. It is the cleanest place to extend the system to new formats."),
            ("List of metrics",
             "All benchmark results land in MetricsCollector.items as FileMetric dataclasses. "
             "Sorting that list by savings_pct produced the per-file chart."),
            ("Set — incremental change detection",
             "I designed the diff that returns four sets: unchanged, recompressed, added, "
             "removed. Set algebra makes the logic obvious."),
        ],
        hardest=(
            "Calibrating the JPG quality presets so that each step gives a meaningful saving "
            "without visibly degrading the image. I built a small reference grid of test images "
            "(gradients, panels, text) and inspected them at q=90, 75, and 55. q=55 was the "
            "first level where the gradients started showing banding, so I pinned that as the "
            "Strong preset. The PNG palette path was also tricky — converting RGBA to P loses "
            "transparency on certain images, so I gated palette reduction on the image mode."
        ),
        learned=[
            "Why image compression has to be lossy and how to expose the trade-off cleanly to the user.",
            "How Pillow's optimize and progressive flags affect the final size.",
            "How to capture honest benchmark numbers — same files, same hardware, same run.",
            "Producing publication-quality charts with matplotlib instead of screenshotting the GUI.",
        ],
        next_time=[
            "Add a side-by-side image preview in the Compress tab so the user sees the quality trade-off live.",
            "Try a dedicated PNG optimiser (e.g. zopfli) for the Strong preset.",
            "Generate a larger benchmark batch with mixed real-world webpages.",
        ],
    )


# ---------------------------------------------------------------
def main() -> None:
    out_paths = [
        build_group_report(),
        build_amir(),
        build_alen(),
        build_diego(),
    ]
    print("Generated:")
    for p in out_paths:
        print(f"  {p}")


if __name__ == "__main__":
    main()
